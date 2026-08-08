import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_resume(resume_data: dict, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = docx.Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    p_info = resume_data.get("personal_info", {})

    # Name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(p_info.get("name", "Candidate").upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Contact
    contact_bits = []
    if p_info.get("email"): contact_bits.append(p_info["email"])
    if p_info.get("phone"): contact_bits.append(p_info["phone"])
    
    loc_raw = p_info.get("location") or ""
    loc_clean = loc_raw.replace(", None", "").replace("None", "").strip(", ")
    if loc_clean: contact_bits.append(loc_clean)

    if p_info.get("linkedin"): contact_bits.append(f"LinkedIn: {p_info['linkedin']}")
    if p_info.get("github"): contact_bits.append(f"GitHub: {p_info['github']}")
    if p_info.get("portfolio"): contact_bits.append(f"Portfolio: {p_info['portfolio']}")

    c_p = doc.add_paragraph()
    c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = c_p.add_run(" | ".join(contact_bits))
    c_run.font.name = 'Calibri'
    c_run.font.size = Pt(9.5)
    c_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Helper function for section headings
    def add_section_heading(text: str):
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(text.upper())
        h_run.font.name = 'Calibri'
        h_run.font.size = Pt(12)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Summary
    if resume_data.get("summary"):
        add_section_heading("Professional Summary")
        s_p = doc.add_paragraph(resume_data["summary"])
        s_p.runs[0].font.name = 'Calibri'
        s_p.runs[0].font.size = Pt(10)

    # Skills
    if resume_data.get("skills"):
        add_section_heading("Technical Skills")
        skills_dict = resume_data["skills"]
        if isinstance(skills_dict, dict):
            for cat, s_list in skills_dict.items():
                if s_list and isinstance(s_list, list):
                    sk_p = doc.add_paragraph()
                    r1 = sk_p.add_run(f"{cat}: ")
                    r1.bold = True
                    r1.font.name = 'Calibri'
                    r1.font.size = Pt(10)
                    r2 = sk_p.add_run(", ".join(s_list))
                    r2.font.name = 'Calibri'
                    r2.font.size = Pt(10)

    # Experience
    if resume_data.get("experiences"):
        add_section_heading("Professional Experience")
        for exp in resume_data["experiences"]:
            ex_p = doc.add_paragraph()
            r1 = ex_p.add_run(exp.get("job_title", "Software Engineer"))
            r1.bold = True
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10.5)
            
            r2 = ex_p.add_run(f" — {exp.get('company', '')} ({exp.get('dates', '')})")
            r2.italic = True
            r2.font.name = 'Calibri'
            r2.font.size = Pt(10)

            bullets = exp.get("bullets", [])
            if not bullets and exp.get("description"):
                bullets = [exp["description"]]
            for bullet in bullets:
                bp = doc.add_paragraph(style='List Bullet')
                brun = bp.add_run(bullet)
                brun.font.name = 'Calibri'
                brun.font.size = Pt(9.5)

    # Projects
    if resume_data.get("projects"):
        add_section_heading("Key Technical Projects")
        for proj in resume_data["projects"]:
            pr_p = doc.add_paragraph()
            r1 = pr_p.add_run(proj.get("name", "Project"))
            r1.bold = True
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10.5)
            
            r2 = pr_p.add_run(f" | {proj.get('technologies', '')}")
            r2.font.name = 'Calibri'
            r2.font.size = Pt(9.5)

            bullets = proj.get("bullets", [])
            if not bullets and proj.get("description"):
                bullets = [proj["description"]]
            for bullet in bullets:
                bp = doc.add_paragraph(style='List Bullet')
                brun = bp.add_run(bullet)
                brun.font.name = 'Calibri'
                brun.font.size = Pt(9.5)

    # Education
    if resume_data.get("education"):
        add_section_heading("Education")
        for edu in resume_data["education"]:
            ed_p = doc.add_paragraph()
            r1 = ed_p.add_run(f"{edu.get('degree', '')} in {edu.get('specialization', '')}")
            r1.bold = True
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10)
            
            r2 = ed_p.add_run(f" — {edu.get('university', '')} ({edu.get('dates', '')})")
            r2.font.name = 'Calibri'
            r2.font.size = Pt(10)

    doc.save(output_path)
    return output_path
